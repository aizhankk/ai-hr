from __future__ import annotations

import io
import os
import re
from datetime import datetime
from typing import Dict, List, Tuple

from docx import Document
from pypdf import PdfReader
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

SKILL_KEYWORDS = [
    "python",
    "java",
    "javascript",
    "typescript",
    "sql",
    "postgresql",
    "mysql",
    "mongodb",
    "redis",
    "fastapi",
    "django",
    "flask",
    "spring",
    "react",
    "node",
    "docker",
    "kubernetes",
    "aws",
    "gcp",
    "azure",
    "machine learning",
    "nlp",
    "pandas",
    "numpy",
    "scikit-learn",
    "tensorflow",
    "pytorch",
]


def _extract_pdf_text(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    chunks: List[str] = []
    for page in reader.pages:
        # Try regular extraction first, then layout mode for PDFs
        # where text ordering is fragmented.
        base_text = page.extract_text() or ""
        if len(base_text.strip()) < 40:
            try:
                base_text = page.extract_text(extraction_mode="layout") or base_text
            except TypeError:
                pass
        chunks.append(base_text)
    return "\n".join(chunks)


def _extract_docx_text(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    chunks: List[str] = [paragraph.text for paragraph in doc.paragraphs]
    # Include table cells, many CV templates keep key data in tables.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.append(cell.text)
    return "\n".join(chunks)


def extract_resume_text(filename: str, file_bytes: bytes) -> str:
    lowered = filename.lower()
    if lowered.endswith(".pdf"):
        return _extract_pdf_text(file_bytes)
    if lowered.endswith(".docx"):
        return _extract_docx_text(file_bytes)
    raw_text = file_bytes.decode("utf-8", errors="ignore")
    return re.sub(r"\s+", " ", raw_text).strip()


def normalize_text(text: str) -> str:
    cleaned = text.replace("\x00", " ")
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    return cleaned


def _refine_text_with_chatgpt(text: str) -> str:
    api_key = os.getenv("OPENAI_API_KEY", "").strip()
    if not api_key or len(text.strip()) < 50:
        return text

    model = os.getenv("OPENAI_TEXT_EXTRACTION_MODEL", "gpt-4o-mini")
    try:
        from openai import OpenAI
    except Exception:
        return text

    prompt = (
        "You are extracting clean plain text from a resume OCR dump. "
        "Fix broken line wraps, remove duplicated headers/footers, keep all important facts, "
        "keep original language, and return only clean resume text without markdown."
    )

    try:
        client = OpenAI(api_key=api_key)
        response = client.responses.create(
            model=model,
            input=[
                {"role": "system", "content": prompt},
                {"role": "user", "content": text[:12000]},
            ],
            temperature=0.0,
        )
        output = getattr(response, "output_text", "") or ""
        return normalize_text(output) if output.strip() else text
    except Exception:
        return text


def detect_skills(text: str) -> List[str]:
    normalized = text.lower()
    found = []
    for skill in SKILL_KEYWORDS:
        # Экранируем спецсимволы (c++, scikit-learn и т.п.)
        pattern = r"(?<![a-z0-9])" + re.escape(skill) + r"(?![a-z0-9])"
        if re.search(pattern, normalized):
            found.append(skill)
    return sorted(set(found))


def detect_years_experience(text: str) -> int:
    lowered = text.lower()
    patterns = [
        r"(\d+)\+?\s*(?:years|year|yrs|yr|лет|года|год)\s*(?:of)?\s*(?:experience|опыта|стажа)?",
        r"(?:experience|опыт|стаж)\s*[:\-]?\s*(\d+)\+?\s*(?:years|year|yrs|yr|лет|года|год)",
    ]
    candidates: List[int] = []
    for pattern in patterns:
        for match in re.findall(pattern, lowered):
            try:
                candidates.append(int(match))
            except ValueError:
                continue
    direct_years = max(candidates) if candidates else 0

    # Fallback: infer from date ranges like 2019-2023 / 2021-present.
    current_year = datetime.utcnow().year
    month_token = r"(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec|янв|фев|мар|апр|май|июн|июл|авг|сен|сент|окт|ноя|дек)\.?"
    range_patterns = [
        r"(20\d{2})\s*[-–]\s*(20\d{2})",
        r"(20\d{2})\s*[-–]\s*(?:present|current|now|настоящее время|по н\.в\.)",
        rf"{month_token}\s*(20\d{{2}})\s*[-–]\s*{month_token}\s*(20\d{{2}})",
        rf"{month_token}\s*(20\d{{2}})\s*[-–]\s*(?:present|current|now|настоящее время|по н\.в\.)",
    ]
    inferred = 0
    for pattern in range_patterns:
        for match in re.findall(pattern, lowered):
            start = int(match[0])
            if len(match) > 1 and match[1].isdigit():
                end = int(match[1])
            else:
                end = current_year
            if 1980 <= start <= current_year and 1980 <= end <= current_year and end >= start:
                inferred = max(inferred, end - start)

    return max(direct_years, inferred)


def detect_education(text: str) -> str:
    lowered = text.lower()
    education_patterns = [
        ("PhD", [r"\bphd\b", r"\bdoctorate\b", r"\bкандидат наук\b", r"\bдоктор наук\b"]),
        ("Master's", [r"\bmaster", r"\bm\.?\s*sc\b", r"\bmsc\b", r"\bмагистр"]),
        ("Bachelor's", [r"\bbachelor", r"\bb\.?\s*sc\b", r"\bbs\b", r"\bбакалавр"]),
        ("College", [r"\bcollege\b", r"\bassociate\b", r"\bколледж"]),
    ]
    for label, patterns in education_patterns:
        for pattern in patterns:
            if re.search(pattern, lowered):
                return label
    return "Not detected"


def rank_resumes_against_job(job_description: str, files: List[Tuple[str, bytes]]) -> List[Dict]:
    normalized_job = normalize_text(job_description)
    docs: List[str] = [normalized_job]
    extracted: List[Dict] = []
    job_skills = set(detect_skills(normalized_job))
    job_tokens = set(re.findall(r"[a-zA-Zа-яА-Я0-9\+\#\.]{3,}", normalized_job.lower()))

    for filename, raw in files:
        text = normalize_text(extract_resume_text(filename, raw))
        text = _refine_text_with_chatgpt(text)
        extracted.append(
            {
                "filename": filename,
                "text": text,
                "skills": detect_skills(text),
                "experience_years": detect_years_experience(text),
                "education": detect_education(text),
            }
        )
        docs.append(text or "")

    # Word-level similarity + char n-gram similarity for better robustness
    # on multilingual text and noisy PDF extraction.
    word_vectorizer = TfidfVectorizer(ngram_range=(1, 2), min_df=1)
    word_matrix = word_vectorizer.fit_transform(docs)
    word_sim = cosine_similarity(word_matrix[0:1], word_matrix[1:]).flatten()

    char_vectorizer = TfidfVectorizer(analyzer="char_wb", ngram_range=(3, 5), min_df=1)
    char_matrix = char_vectorizer.fit_transform(docs)
    char_sim = cosine_similarity(char_matrix[0:1], char_matrix[1:]).flatten()

    semantic_similarities = 0.65 * word_sim + 0.35 * char_sim

    ranked: List[Dict] = []
    for idx, item in enumerate(extracted):
        semantic_score = float(max(0.0, min(1.0, semantic_similarities[idx])))

        resume_skills = set(item["skills"])
        if job_skills:
            skill_overlap = len(job_skills & resume_skills) / len(job_skills)
        else:
            skill_overlap = 0.0

        resume_tokens = set(re.findall(r"[a-zA-Zа-яА-Я0-9\+\#\.]{3,}", item["text"].lower()))
        if job_tokens:
            keyword_coverage = len(job_tokens & resume_tokens) / len(job_tokens)
        else:
            keyword_coverage = 0.0

        # Recruiter-facing score:
        # semantic relevance + explicit skill match + keyword coverage.
        raw_score = (
            0.45 * semantic_score
            + 0.40 * skill_overlap
            + 0.15 * keyword_coverage
        )
        # Guardrail: if required skills exist but overlap is low, down-weight the total.
        # This prevents cases like 0/N required skills still showing ~50% only from semantics.
        if job_skills:
            if skill_overlap <= 0.0:
                # Hard rule: no required-skill match -> not a fit.
                raw_score = 0.0
            else:
                # Keep gating, but make it softer to avoid under-scoring good profiles.
                skill_gate = 0.82 + 0.18 * skill_overlap
                raw_score *= skill_gate
                # Bonus for strong/very strong required-skill match.
                if skill_overlap >= 0.60:
                    raw_score = min(1.0, raw_score + 0.05)
                if skill_overlap >= 0.80:
                    raw_score = min(1.0, raw_score + 0.07)
        score = float(round(max(0.0, min(1.0, raw_score)) * 100, 2))
        # Recruiter-friendly floors for candidates with meaningful must-have coverage.
        if job_skills and skill_overlap >= 0.80 and score < 75.0:
            score = 75.0
        elif job_skills and skill_overlap >= 0.60 and score < 65.0:
            score = 65.0
        ranked.append(
            {
                "rank": 0,
                "candidate_name": item["filename"].rsplit(".", 1)[0],
                "filename": item["filename"],
                "match_score": score,
                "semantic_score": round(semantic_score * 100, 2),
                "skill_overlap": round(skill_overlap * 100, 2),
                "keyword_coverage": round(keyword_coverage * 100, 2),
                "experience_years": item["experience_years"],
                "education": item["education"],
                "skills": item["skills"][:8],
                "preview": item["text"][:600],
            }
        )

    ranked.sort(key=lambda row: row["match_score"], reverse=True)
    for i, row in enumerate(ranked, start=1):
        row["rank"] = i
    return ranked
